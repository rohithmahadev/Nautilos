import streamlit as st 
import pandas as pd 
from langchain_community.document_loaders import PyPDFLoader
import tempfile
import os
from dotenv import load_dotenv 
import pandas as pd


from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


from groq import Groq
from langchain_groq import ChatGroq

from langchain_community.document_loaders import WebBaseLoader
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_community.document_loaders import PyPDFLoader
from utils import parsing_pdf, accomplishments_pdf 
load_dotenv()


os.getenv("GROQ_API_KEY")
llm = ChatGroq(
    model="llama-3.1-8b-instant",
    temperature=0,
   
)

uploaded_pdf_file = st.file_uploader("Upload a PDF version of your resume",
                                     type= "pdf",
                                     accept_multiple_files= False)
        
st.session_state.pdf_data = parsing_pdf(uploaded_pdf_file)





uploaded_accomplishments_pdf_file = st.file_uploader("Upload a PDF version of your accomplishments",
                                     type= "pdf",
                                     accept_multiple_files= False)


st.session_state.accomplishment_data = accomplishments_pdf(uploaded_accomplishments_pdf_file)



st.header("Enter the job description link that you want to tailor your resume")

st.markdown("Note: Must be open source")

website = st.text_input("Link: ")



### LLM 

prompt_extract_career = PromptTemplate.from_template(
 """
 ##Job Description
 {page_data}
 ### Instruction:
 The scrapted text is from the career's page of a company's website. Look in detail for the skills. It might be in other headings too. The instruction for you is to extract the relavent skills from the job description. 

 Things to note:

 1. Remember only the skills needed to extracted. 
 2. Return all the skills
 3. Look for main keywords and identify the soft skills needed too.
 4. Look for cross funcitional skills

Only return the valid JSON file.

### Valid JSON Format without any preamble
 
 """ 
)


prompt_tailor_resume = PromptTemplate.from_template(
"""
### Job Requirements
{job_data}

### Current Resume
{resume_data}

### Accomplishments Bank
{accomplishments_data}

### Instruction
You are a professional resume writer. Create a tailored one-page resume optimized for this job.

Rules:
1. Use ONLY information from the resume and accomplishments - do NOT hallucinate
2. Prioritize experience/skills that match the job requirements
3. Reword bullets to include keywords from the job description
4. Add relevant accomplishments from the accomplishments bank if they strengthen the application
5. Keep the most impactful 4-6 bullets per job/section
6. If adding accomplishments, put them in the appropriate section (Experience, Projects, or Personal Accomplishments)
7. When extracting skills, extract only from the job data, resume data and match them,drop the unmatched or unwanted skills. Drop based on relavency to the job data.

Sample JSON format to return (Add extra heading if mentioned in the resume for exampel research):
{{{{
  "professional_summary": "2-3 sentences highlighting fit for this role",
  "skills": ["skill1", "skill2","skill3","skill4".....],
  "experience": [
    {{{{
      "company": "Company name",
      "role": "Job title",
      "duration": "Date range",
      "bullets": ["bullet1", "bullet2"]
    }}}}
  ],
  "projects": [
    {{{{
      "project_name": "project name",
      "project_description": "project description"
    }}}}
  ],
  "education": "Education details",
  "certifications": ["cert1", "cert2"]
}}}}

### CRITICAL: Return ONLY the JSON object. No markdown, no code blocks, no explanations.
"""
)        


def generate_resume_docx(tailored_data):
    """Generate a professional resume Word document from tailored data"""
    
    doc = Document()
    
    # Professional Summary
    if 'professional_summary' in tailored_data:
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(tailored_data['professional_summary'])
    
    # Skills
    if 'skills' in tailored_data and tailored_data['skills']:
        doc.add_heading('Skills', level=1)
        skills_text = ', '.join(tailored_data['skills'])
        doc.add_paragraph(skills_text)
    
    # Experience
    if 'experience' in tailored_data and tailored_data['experience']:
        doc.add_heading('Professional Experience', level=1)
        for exp in tailored_data['experience']:
            # Company and role
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('role', '')} | {exp.get('company', '')}").bold = True
            
            # Duration
            if 'duration' in exp:
                doc.add_paragraph(exp['duration'])
            
            # Bullets
            if 'bullets' in exp:
                for bullet in exp['bullets']:
                    doc.add_paragraph(bullet, style='List Bullet')
            
            doc.add_paragraph()  # Spacing
    
    # Projects
    if 'projects' in tailored_data and tailored_data['projects']:
        doc.add_heading('Projects', level=1)
        for project in tailored_data['projects']:
            # Project name (bold)
            p = doc.add_paragraph()
            p.add_run(project.get('project_name', '')).bold = True
            
            # Project description
            doc.add_paragraph(project.get('project_description', ''))
            doc.add_paragraph()  # Spacing
    
    # Education
    if 'education' in tailored_data and tailored_data['education']:
        doc.add_heading('Education', level=1)
        doc.add_paragraph(tailored_data['education'])
    
    # Certifications
    if 'certifications' in tailored_data and tailored_data['certifications']:
        doc.add_heading('Certifications', level=1)
        for cert in tailored_data['certifications']:
            doc.add_paragraph(cert, style='List Bullet')
    
    # Save to BytesIO for Streamlit download
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

json_parser = JsonOutputParser()
# Inside the button click
if st.button("Submit"):
    # Validation checks
    if website == "" or website is None:
        st.error("⚠️ Please enter a job description link to proceed!")
    elif uploaded_pdf_file is None:
        st.error("⚠️ Please upload your resume PDF to proceed!")
    else:
        # Extract job data
        with st.spinner("Loading job data..."):
            loader = WebBaseLoader(website)
            st.session_state.career_page_data = loader.load()
            chain_extract = prompt_extract_career | llm
            res = chain_extract.invoke(input={'page_data': st.session_state.career_page_data})
            st.session_state.json_parsed_job_data = json_parser.parse(res.content)
            
            display_data = st.session_state.json_parsed_job_data.copy()
            st.write(display_data)
            df = pd.DataFrame([display_data]).T
            df.columns = ['Job Description Details']
            st.table(df)
        
        st.success("✅ Job data loaded successfully!")
        
        # Tailor the resume (INSIDE the button block)
        with st.spinner("Tailoring your resume..."):
            chain_tailor = prompt_tailor_resume | llm
            res_tailor = chain_tailor.invoke(input={
                'job_data': st.session_state.json_parsed_job_data,
                'resume_data': st.session_state.pdf_data,
                'accomplishments_data': st.session_state.accomplishment_data
            })
            
            st.session_state.tailored_resume = json_parser.parse(res_tailor.content)
            st.success("✅ Resume tailored successfully!")
        
        # Generate Word document (STILL INSIDE the button block)
        doc_file = generate_resume_docx(st.session_state.tailored_resume)
        
        # Download button
        st.download_button(
            label="📥 Download Tailored Resume",
            data=doc_file,
            file_name="tailored_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        st.write("The output is in word format and use it as a reference to edit your resume.")


